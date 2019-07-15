import os
import sys
import string
from sets import Set
try:
	from docx import Document
except:
	os.system("pip install python-docx")
#from docx.oxml import nsdecls, parse_xml
from docx.shared import Inches, Cm
from docx.enum.text import WD_BREAK
from docx.table import Table
import time, datetime


def remove_dupes(jaja):
	uniqueList = Set(jaja)
	return uniqueList

def cp_collect(): #copy and paste a list of items and it will kick it out as a unique list
	master = []
	while 1:
		host = raw_input("  > ")
		if host == "done":
			unique_master = remove_dupes(master)
			return unique_master
		master.append(host)
	unique_master = remove_dupes(master)
	return unique_master


def cp_text(): #copy and paste a list of items and it will kick it out as a unique list
	master = []
	while 1:
		line = raw_input("txt> ")
		if line == "done":
			master_string = ''.join(master)
			return master_string
		master.append(line+"\n")
	master_string = ''.join(master)
	return master_string


class Pub:
    document = ""
    ProjectName = ""
    # Create (and open) the DOCX document
    def __init__(self):
        print "Ready for publication..."
    #changing the page margins

    def create_new_report(self):
        self.ProjectName = raw_input("Client Name: ")
        print "Initializing new report for", self.ProjectName
        self.document = Document()
        sections = self.document.sections
        for section in sections:
            section.top_margin      = Cm(0.5)
            section.bottom_margin   = Cm(0.5)
            section.left_margin     = Cm(1.0)
            section.right_margin    = Cm(1.0)

    def store_new_report(self, pName):
        self.ProjectName = pName
        print "Initializing new report for", self.ProjectName
        self.document = Document()
        sections = self.document.sections
        for section in sections:
            section.top_margin      = Cm(0.5)
            section.bottom_margin   = Cm(0.5)
            section.left_margin     = Cm(1.0)
            section.right_margin    = Cm(1.0)


    def create_ip_entry(self):
        self.document.add_heading("Address: " + IP, 0)
        self.document.add_paragraph("Hostname: " + HOSTNAME)
        self.document.add_paragraph("FQDN: ", + FQDN)
        self.document.add_heading("Open Ports and Services", level=1)
        self.document.add_paragraph("<PORTS>", style='IntenseQuote')
        self.document.add_heading("Attack Surface Investigation", level=1)
        self.document.add_paragraph("<INVESTIGATION>")
        self.document.add_heading("Findings", level=1)

        self.document.add_paragraph("Finding1", style='ListBullet')
        self.document.add_paragraph("Finding2", style='ListBullet')
        self.document.add_heading("Recommendations", level=1)
        self.document.add_paragraph("<Recommendations>")

    def create_vuln_entry(self):
        FindingVal = raw_input("Finding: ")
        DescriptionVal = raw_input("Description: ")
        ExploitableVal = raw_input("Exploit Research: ")
        print "Affected Hosts: "
        AffectedHosts = cp_collect()
        #for item in AffectedHosts:
        #        self.document.add_paragraph(item, style='ListBullet')
        RecommendationsVal = raw_input("Recommendations: ")

        print "Finding:          ", FindingVal
        print "Description:      ", DescriptionVal
        print "Exploit Research: ", ExploitableVal
        print "Affected hosts:   "
        for item in AffectedHosts:
                print "    [+] "+item
        print "Recommendations: ", RecommendationsVal
                #self.document.add_paragraph(item, style='ListBullet')
        print "--------------"
        contin = raw_input()

        FINDING         = FindingVal
        DESCRIPTION     = DescriptionVal
        EXPLOITATION    = ExploitableVal
        RECOMMENDATIONS = RecommendationsVal

        self.document.add_heading(FINDING, 0)
        self.document.add_heading("Description:", level=1)
        self.document.add_paragraph(DESCRIPTION)#, style='IntenseQuote')
        self.document.add_heading("Exploit Research", level=1)
        self.document.add_paragraph(EXPLOITATION)
        self.document.add_heading("Affected Hosts", level=1)
        for item in AffectedHosts:
            self.document.add_paragraph(item, style='ListBullet')
        self.document.add_heading("Recommendations", level=1)
        self.document.add_paragraph(RECOMMENDATIONS)
        #self.document.add_paragraph("", style='IntenseQuote')       # small DIV
        self.document.add_paragraph("")

    # Generate the skeleton of the report
    def save_report(self):
        ts = time.time()
        timestamp = datetime.datetime.fromtimestamp(ts).strftime('%Y%m%d')#-%H-%M-%S')
        print "Saving report at", timestamp
        pName = self.ProjectName.replace(" ", "_")
        #self.document.save('pentest_demo.docx')
        filename = pName+"_Pubsy_"+timestamp+".docx"
        self.document.save(filename)
        print "New assessment report created:", filename

    def run_command(self):
        cmd = raw_input("cmd> ")
        process = os.popen(cmd)
        preprocessed = process.read()
        process.close()
        table = self.document.add_table(rows=1, cols=1, style='TableGrid')
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = preprocessed
        self.document.add_paragraph("")

    def store_command(self, cmdText):
        cmd = cmdText
        process = os.popen(cmd)
        preprocessed = process.read()
        process.close()
        table = self.document.add_table(rows=1, cols=1, style='TableGrid')
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = preprocessed
        self.document.add_paragraph("")

    def output_content(self):
        #shading_red_1 = docx.oxml.parse_xml(r'<w:shd {} w:fill="FFFFFF"/>'.format(docx.oxml.ns.nsdecls('w')))
        print "Enter Text Output: "
        output = cp_text()              #cp_text allows you to copy text in, which converts the cp'd list to a string
        table = self.document.add_table(rows=1, cols=1, style='TableGrid')
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = output
        self.document.add_paragraph("")


    def pub_header(self, text="<Section Header>"):
        self.document.add_heading(text, 0)

    def pub_subheader(self, text):
        self.document.add_heading(text, 1)

    def pub_paragraph(self, title="<Paragragh Title>"):
        pub_text = raw_input("c&p> ")
        self.document.add_heading(title, level=1)
        self.document.add_paragraph(pub_text)
    
    def store_paragraph(self, pub_text):
        #self.document.add_heading(title, level=1)
        self.document.add_paragraph(pub_text)

    def pub_bullets(self):
        bullets = cp_text()
        for item in bullets:
            self.document.add_paragraph(item, style='ListBullet')

    def store_textbox(self, boxText):
        text = boxText
        table = self.document.add_table(rows=1, cols=1, style='TableGrid')
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = text
        self.document.add_paragraph("")


    def pub_textbox(self):
        text = cp_text()
        table = self.document.add_table(rows=1, cols=1, style='TableGrid')
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = text
        self.document.add_paragraph("")

    def pub_bold_header(self):
        key = raw_input("Bold key: ")
        value = raw_input("Normal Value: ")
        p = self.document.add_paragraph(" ")
        p.add_run(key+": ").bold = True
        p.add_run(value)

# COMMAND LINE INTERFACE
def console():
	os.system("cls")
	print "\n\n"

	print "  ____      _   _    ____    ____      __   __   "
	print "U|  _\"\ uU |\"|u| |U | __\")u / __\"| u   \ \ / /   "
	print "\| |_) |/ \| |\| | \|  _ \/<\___ \/     \ V /    "
	print " |  __/    | |_| |  | |_) | u___) |    U_|\"|_u   "
	print " |_|      <<\___/   |____/  |____/>>     |_|     "
	print " ||>>_   (__) )(   _|| \\_   )(  (__).-,//|(_    "
	print " (__)__)      (__) (__) (__) (__)      \_) (__)  "
	print "*****************************************************"
	print r"*  Pubsy               Pentest Report Publication  *"
	print r"*  Console                      Version 1.5        *"   
	print r"****************************************************" 

	print " Type 'help'\n"
	cmd = ''
	report = Pub()

	while (string.upper(cmd)) != ("QUIT" or "EXIT"):
		cmd = raw_input("\nPubsy#> ")
		command = string.split(cmd, " ")

		#Informational Commands
		if string.upper(command[0]) == "HELP" and len(command) <= 1:
			print __helpfile__
		elif string.upper(command[0]) == "NEW":
			report.create_new_report()
		elif string.upper(command[0]) == "VULN":						#> assets
			report.create_vuln_entry()
		elif string.upper(command[0]) == "IP":
			report.create_ip_entry()
		elif string.upper(command[0]) == "OUTPUT":
			report.output_content()
		elif string.upper(command[0]) == "RUN":
			report.run_command()
		elif string.upper(command[0]) == "SAVE":
			report.save_report()
		elif string.upper(command[0]) == "HEADER":
			try:
				report.pub_header(' '.join(command[1:]))
			except:
				report.pub_header()
		elif string.upper(command[0]) == "SUBHEADER":
			try:
				report.pub_subheader(' '.join(command[1:]))
			except:
				report.pub_subheader()                        
		elif string.upper(command[0]) == "PARAGRAPH":
			try:
				report.pub_paragraph(' '.join(command[1:]))
			except:
				report.pub_paragraph()
				
		elif string.upper(command[0]) == "BULLETS":
			report.pub_bullets()

		elif string.upper(command[0]) == "BOLDSUB":
			report.pub_bold_header()
			
		elif string.upper(command[0]) == "TEXTBOX":
			report.pub_paragraph()                     
	# Catch all
		else:
			pass
		


__helpfile__='''
Commands
--------
[+] Pubsy Console v1.0 - Dynamically Create Quick Pentest Reports from Console Output

new                            create a new pentest report
view                           open and view the pentest report
titlepage                      create a title page
vuln                            publish a vulnerability report entry
ip                              publish an ip report entry
output                          
run
save

'''


    #ip = {"Address":"",
            #"Hostname":"",
            #"FQDN":"",
            #"Open Ports and Services":[],
            #"Attack Surface Investigation":"",
            #"Findings":[],
            #"Recommendations":""}
#new_finding                    creates a new Vuln Report Entry requiring additional detail
    #vuln = {"Finding":"",
            #"Description":"",
            #"Affected Hosts":[],
            #"Details":[],
            #"Severity Level":"",
#conclusion
#assets                          view existing assets
#'''

if __name__=='__main__':
    console()
